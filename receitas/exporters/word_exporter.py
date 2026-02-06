"""Word exporter"""

from docx import Document
from docx.shared import Pt


def exportar_word(receita, destino):
    doc = Document()

    # Título
    titulo = doc.add_heading(receita.titulo, level=0)
    titulo.runs[0].font.size = Pt(20)

    # Fonte
    p = doc.add_paragraph()
    p.add_run(f"Fonte: {receita.fonte}\n").italic = True
    p.add_run(f"URL: {receita.url}").italic = True

    # Ingredientes
    doc.add_heading("Ingredientes", level=1)
    for ing in receita.ingredientes:
        doc.add_paragraph(ing, style="List Bullet")

    # Preparo
    doc.add_heading("Modo de preparo", level=1)
    for passo in receita.preparo:
        doc.add_paragraph(passo, style="List Number")

    doc.save(destino)
